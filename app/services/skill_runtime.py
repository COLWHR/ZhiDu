from __future__ import annotations

import csv
import hashlib
import io
import json
import logging
import os
import re
import textwrap
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

logger = logging.getLogger(__name__)


PROJECT_ROOT = Path(__file__).resolve().parents[2]
UPLOADS_ROOT = PROJECT_ROOT / "uploads"
SKILL_REPOSITORY_ROOT = PROJECT_ROOT / "skills"
UPLOADS_ROOT.mkdir(parents=True, exist_ok=True)
(UPLOADS_ROOT / "attachments").mkdir(parents=True, exist_ok=True)
(UPLOADS_ROOT / "artifacts").mkdir(parents=True, exist_ok=True)


DEFAULT_SKILL_CATALOG: list[dict[str, Any]] = [
    {
        "skill_key": "chat.reply",
        "name": "智能对话",
        "category": "conversation",
        "description": "基础聊天与问答能力。",
        "input_modalities": ["text"],
        "output_types": ["text"],
        "required_models": ["fast", "reasoning"],
        "required_tools": [],
        "params_schema": {},
        "permission_scope": ["chat"],
        "cost_level": "low",
        "status": "active",
    },
    {
        "skill_key": "image.view",
        "name": "图片理解",
        "category": "multimodal",
        "description": "理解图片内容并基于图片回答。",
        "input_modalities": ["image", "text"],
        "output_types": ["text"],
        "required_models": ["vision"],
        "required_tools": [],
        "params_schema": {},
        "permission_scope": ["read_media"],
        "cost_level": "medium",
        "status": "active",
    },
    {
        "skill_key": "image.generate",
        "name": "图片生成",
        "category": "generation",
        "description": "根据描述生成图片或插图。",
        "input_modalities": ["text", "image"],
        "output_types": ["image", "artifact"],
        "required_models": ["reasoning"],
        "required_tools": ["image_renderer"],
        "params_schema": {},
        "permission_scope": ["write_artifact"],
        "cost_level": "high",
        "status": "active",
    },
    {
        "skill_key": "file.read",
        "name": "文件读取",
        "category": "document",
        "description": "读取文本、docx、csv、json 等文件内容。",
        "input_modalities": ["file"],
        "output_types": ["text"],
        "required_models": ["fast"],
        "required_tools": ["file_parser"],
        "params_schema": {},
        "permission_scope": ["read_file"],
        "cost_level": "low",
        "status": "active",
    },
    {
        "skill_key": "file.summary",
        "name": "文件总结",
        "category": "document",
        "description": "对上传文件进行摘要、归纳与提炼。",
        "input_modalities": ["file", "text"],
        "output_types": ["text"],
        "required_models": ["fast", "reasoning"],
        "required_tools": ["file_parser"],
        "params_schema": {},
        "permission_scope": ["read_file"],
        "cost_level": "medium",
        "status": "active",
    },
    {
        "skill_key": "doc.generate",
        "name": "文档生成",
        "category": "generation",
        "description": "生成可下载的文档成果物。",
        "input_modalities": ["text", "file"],
        "output_types": ["artifact", "file"],
        "required_models": ["reasoning"],
        "required_tools": ["docx_renderer"],
        "params_schema": {},
        "permission_scope": ["write_artifact"],
        "cost_level": "high",
        "status": "active",
    },
    {
        "skill_key": "sheet.generate",
        "name": "表格生成",
        "category": "generation",
        "description": "生成结构化数据表格或 CSV 成果物。",
        "input_modalities": ["text", "file"],
        "output_types": ["artifact", "file"],
        "required_models": ["reasoning"],
        "required_tools": ["table_renderer"],
        "params_schema": {},
        "permission_scope": ["write_artifact"],
        "cost_level": "medium",
        "status": "active",
    },
    {
        "skill_key": "audio.transcribe",
        "name": "语音转写",
        "category": "multimodal",
        "description": "将音频转写成文字并继续处理。",
        "input_modalities": ["audio"],
        "output_types": ["text"],
        "required_models": ["transcribe"],
        "required_tools": ["asr"],
        "params_schema": {},
        "permission_scope": ["read_media"],
        "cost_level": "medium",
        "status": "active",
    },
    {
        "skill_key": "video.analyze",
        "name": "视频理解",
        "category": "multimodal",
        "description": "分析视频内容、关键帧和字幕。",
        "input_modalities": ["video"],
        "output_types": ["text"],
        "required_models": ["vision", "reasoning"],
        "required_tools": ["video_sampler"],
        "params_schema": {},
        "permission_scope": ["read_media"],
        "cost_level": "high",
        "status": "active",
    },
    {
        "skill_key": "web.search",
        "name": "联网搜索",
        "category": "tooling",
        "description": "根据需求检索互联网信息。",
        "input_modalities": ["text"],
        "output_types": ["text"],
        "required_models": ["fast"],
        "required_tools": ["web_search"],
        "params_schema": {},
        "permission_scope": ["network"],
        "cost_level": "medium",
        "status": "active",
    },
    {
        "skill_key": "artifact.publish",
        "name": "成果发布",
        "category": "workflow",
        "description": "将生成的成果物整理成可下载资源。",
        "input_modalities": ["text", "artifact"],
        "output_types": ["artifact"],
        "required_models": ["reasoning"],
        "required_tools": ["artifact_store"],
        "params_schema": {},
        "permission_scope": ["write_artifact"],
        "cost_level": "low",
        "status": "active",
    },
]


def skill_lookup_map() -> dict[str, dict[str, Any]]:
    return {item["skill_key"]: item for item in DEFAULT_SKILL_CATALOG}


def get_skill_catalog() -> list[dict[str, Any]]:
    catalog = list(DEFAULT_SKILL_CATALOG)
    try:
        from app.services.skillhub_repository import load_repository_skills

        repository_skills = load_repository_skills(SKILL_REPOSITORY_ROOT)
    except Exception as exc:
        logger.warning("Failed to load local skill repository: %s", exc)
        repository_skills = []

    by_key = {skill["skill_key"]: skill for skill in catalog}
    for skill in repository_skills:
        by_key[skill["skill_key"]] = skill
    return list(by_key.values())


def infer_attachment_kind(mime_type: Optional[str], file_name: str) -> str:
    mime = (mime_type or "").lower()
    name = (file_name or "").lower()
    if mime.startswith("image/") or name.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp")):
        return "image"
    if mime.startswith("video/") or name.endswith((".mp4", ".mov", ".avi", ".webm", ".mkv")):
        return "video"
    if mime.startswith("audio/") or name.endswith((".mp3", ".wav", ".m4a", ".ogg", ".flac")):
        return "audio"
    if mime in {"application/pdf"} or name.endswith(".pdf"):
        return "file"
    if name.endswith((".docx", ".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm", ".py", ".ts", ".js", ".jsonl")):
        return "file"
    return "file"


def is_textual_attachment(file_name: str, mime_type: Optional[str] = None) -> bool:
    kind = infer_attachment_kind(mime_type, file_name)
    if kind != "file":
        return False
    name = (file_name or "").lower()
    return name.endswith((".txt", ".md", ".csv", ".json", ".jsonl", ".xml", ".html", ".htm", ".py", ".ts", ".js", ".css", ".yml", ".yaml"))


def _safe_storage_path(relative_path: str) -> Path:
    relative = Path(relative_path.lstrip("/"))
    return PROJECT_ROOT / relative


def _load_docx_document():
    try:
        from docx import Document

        return Document
    except Exception:
        return None


def _load_pillow_modules():
    try:
        from PIL import Image, ImageDraw, ImageFont

        return Image, ImageDraw, ImageFont
    except Exception:
        return None, None, None


def attachment_public_url(relative_path: Path | str) -> str:
    rel = str(relative_path).replace("\\", "/")
    if rel.startswith("uploads/"):
        return "/" + rel
    if rel.startswith("/uploads/"):
        return rel
    return "/uploads/" + rel.lstrip("/")


def make_attachment_path(file_name: str, prefix: str = "attachments") -> Path:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", file_name).strip("_") or "upload.bin"
    return UPLOADS_ROOT / prefix / f"{stamp}_{safe_name}"


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def read_text_attachment(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".json", ".jsonl", ".csv", ".xml", ".html", ".htm", ".py", ".ts", ".js", ".css", ".yml", ".yaml"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        Document = _load_docx_document()
        if Document is None:
            return ""
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return ""


def summarize_attachment_file(path: Path, mime_type: Optional[str] = None) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".json", ".jsonl", ".csv", ".xml", ".html", ".htm", ".py", ".ts", ".js", ".css", ".yml", ".yaml", ".docx"}:
        text = read_text_attachment(path)
        if not text.strip():
            return ""
        return text[:6000]

    if suffix in {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}:
        try:
            with Image.open(path) as img:
                return f"[image {img.width}x{img.height}] {path.name}"
        except Exception:
            return f"[image] {path.name}"

    if suffix in {".mp4", ".mov", ".avi", ".webm", ".mkv", ".mp3", ".wav", ".m4a", ".ogg", ".flac"}:
        return f"[{infer_attachment_kind(mime_type, path.name)}] {path.name}"

    return f"[file] {path.name}"


def extract_attachment_text(attachment: Dict[str, Any]) -> str:
    storage_url = attachment.get("storage_url") or attachment.get("url") or ""
    if not storage_url:
        return ""
    if storage_url.startswith("http://") or storage_url.startswith("https://"):
        return ""
    path = _safe_storage_path(storage_url)
    if not path.exists():
        return ""
    if path.stat().st_size > 10 * 1024 * 1024:
        return f"[file too large to inline: {path.name}]"
    return summarize_attachment_file(path, attachment.get("mime_type"))


def latest_user_text(context_messages: List[dict[str, Any]]) -> str:
    for msg in reversed(context_messages or []):
        speaker = str(msg.get("speaker") or msg.get("role") or "").strip().lower()
        if speaker in {"user", "用户", "鐢ㄦ埛"}:
            text = str(msg.get("content") or "").strip()
            if text:
                return text
    return ""


def infer_skill_route(persona: Dict[str, Any], context_messages: List[dict[str, Any]], attachments: List[dict[str, Any]]) -> dict[str, Any]:
    skill_keys = set(map(str, persona.get("skills") or ["chat.reply"]))
    text = latest_user_text(context_messages)
    attachment_kinds = [infer_attachment_kind(att.get("mime_type"), att.get("file_name", "")) for att in attachments or []]

    requested_doc = bool(re.search(r"(生成|导出|文档|报告|方案|总结|表格|PPT|演示|文件)", text, re.IGNORECASE))
    if any(kind == "image" for kind in attachment_kinds) and "image.view" in skill_keys:
        return {"skill_key": "image.view", "use_vision": True, "should_generate_artifact": False}
    if any(kind == "audio" for kind in attachment_kinds) and "audio.transcribe" in skill_keys:
        return {"skill_key": "audio.transcribe", "use_vision": False, "should_generate_artifact": False}
    if any(kind == "video" for kind in attachment_kinds) and "video.analyze" in skill_keys:
        return {"skill_key": "video.analyze", "use_vision": False, "should_generate_artifact": False}
    if any(kind == "file" for kind in attachment_kinds) and ("file.read" in skill_keys or "file.summary" in skill_keys):
        return {"skill_key": "file.summary" if "file.summary" in skill_keys else "file.read", "use_vision": False, "should_generate_artifact": False}
    if requested_doc and "doc.generate" in skill_keys:
        return {"skill_key": "doc.generate", "use_vision": False, "should_generate_artifact": True, "artifact_type": "docx"}
    if requested_doc and "sheet.generate" in skill_keys:
        return {"skill_key": "sheet.generate", "use_vision": False, "should_generate_artifact": True, "artifact_type": "csv"}
    return {"skill_key": "chat.reply", "use_vision": any(kind == "image" for kind in attachment_kinds), "should_generate_artifact": False}


def build_llm_messages(
    persona: Dict[str, Any],
    context_messages: List[dict[str, Any]],
    attachments: List[dict[str, Any]],
    base_url: str = "",
) -> tuple[list[dict[str, Any]], bool, dict[str, Any]]:
    route = infer_skill_route(persona, context_messages, attachments)
    use_vision = bool(route.get("use_vision"))
    system_prompt = persona.get("system_prompt") or "你是一个专业的智能体。"
    skill_list = ", ".join(persona.get("skills") or ["chat.reply"])
    system_prompt = f"{system_prompt}\n\n可用技能: {skill_list}"
    messages: list[dict[str, Any]] = [{"role": "system", "content": system_prompt}]

    def _normalize_parts(parts: Iterable[dict[str, Any]]) -> tuple[list[dict[str, Any]], bool]:
        normalized: list[dict[str, Any]] = []
        has_image = False
        for part in parts:
            part_type = str(part.get("type") or "").strip()
            if part_type == "text":
                text = str(part.get("text") or "").strip()
                if text:
                    normalized.append({"type": "text", "text": text})
            elif part_type in {"image", "image_url"}:
                url = str(part.get("url") or part.get("image_url", {}).get("url") or "").strip()
                if url:
                    has_image = True
                    if not url.startswith(("http://", "https://")) and base_url:
                        if url.startswith("/"):
                            url = f"{base_url.rstrip('/')}{url}"
                        else:
                            url = f"{base_url.rstrip('/')}/{url.lstrip('/')}"
                    normalized.append({"type": "image_url", "image_url": {"url": url}})
            elif part_type == "file":
                file_name = str(part.get("file_name") or part.get("name") or "file").strip()
                text = str(part.get("text") or "").strip()
                if text:
                    normalized.append({"type": "text", "text": f"[文件:{file_name}]\n{text}"})
                else:
                    normalized.append({"type": "text", "text": f"[文件:{file_name}]"})
            elif part_type == "video":
                normalized.append({"type": "text", "text": f"[视频:{part.get('file_name') or 'video'}]"})
            elif part_type == "audio":
                normalized.append({"type": "text", "text": f"[音频:{part.get('file_name') or 'audio'}]"})
        return normalized, has_image

    for msg in context_messages or []:
        speaker = str(msg.get("speaker") or msg.get("role") or "").strip()
        role = "user" if speaker in {"用户", "user", "鐢ㄦ埛"} else "assistant"
        parts: list[dict[str, Any]] = []

        if isinstance(msg.get("parts"), list):
            parts.extend(msg["parts"])
        elif isinstance(msg.get("content"), list):
            parts.extend(msg["content"])
        else:
            content_text = str(msg.get("content") or "").strip()
            if content_text:
                parts.append({"type": "text", "text": content_text})

        msg_attachments = msg.get("attachments") or []
        for att in msg_attachments:
            if not isinstance(att, dict):
                continue
            kind = infer_attachment_kind(att.get("mime_type"), att.get("file_name", ""))
            if kind == "image":
                parts.append({"type": "image_url", "url": att.get("storage_url") or att.get("url")})
            else:
                extracted = extract_attachment_text(att)
                if extracted:
                    parts.append({"type": "text", "text": extracted})
                else:
                    parts.append({"type": "text", "text": f"[{kind}:{att.get('file_name') or 'attachment'}]"})

        normalized, has_image = _normalize_parts(parts)
        if has_image:
            use_vision = True
        if not normalized:
            continue
        if len(normalized) == 1 and normalized[0].get("type") == "text":
            messages.append({"role": role, "content": normalized[0]["text"]})
        else:
            messages.append({"role": role, "content": normalized})

    return messages, use_vision, route


def render_text_artifact(text: str, title: str = "artifact") -> tuple[str, str, str]:
    safe_title = re.sub(r"[^A-Za-z0-9._-]+", "_", title).strip("_") or "artifact"
    Document = _load_docx_document()
    if Document is not None:
        try:
            file_name = f"{safe_title}.docx"
            path = UPLOADS_ROOT / "artifacts" / file_name
            doc = Document()
            doc.add_heading(title, level=1)
            for paragraph in text.splitlines():
                doc.add_paragraph(paragraph)
            doc.save(str(path))
            return file_name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", attachment_public_url(path.relative_to(PROJECT_ROOT))
        except Exception as exc:
            logger.warning("Falling back to plain text artifact for %s: %s", title, exc)

    file_name = f"{safe_title}.txt"
    path = UPLOADS_ROOT / "artifacts" / file_name
    content = title + "\n\n" + (text or "")
    path.write_text(content, encoding="utf-8")
    return file_name, "text/plain", attachment_public_url(path.relative_to(PROJECT_ROOT))


def render_csv_artifact(rows: List[Dict[str, Any]], title: str = "artifact") -> tuple[str, str, str]:
    safe_title = re.sub(r"[^A-Za-z0-9._-]+", "_", title).strip("_") or "artifact"
    file_name = f"{safe_title}.csv"
    path = UPLOADS_ROOT / "artifacts" / file_name
    if rows:
        headers = list(rows[0].keys())
    else:
        headers = ["value"]
        rows = [{"value": ""}]
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=headers)
        writer.writeheader()
        writer.writerows(rows)
    return file_name, "text/csv", attachment_public_url(path.relative_to(PROJECT_ROOT))


def render_png_artifact(text: str, title: str = "artifact") -> tuple[str, str, str]:
    safe_title = re.sub(r"[^A-Za-z0-9._-]+", "_", title).strip("_") or "artifact"
    Image, ImageDraw, ImageFont = _load_pillow_modules()
    if Image is None or ImageDraw is None or ImageFont is None:
        return render_text_artifact(f"[PNG fallback]\n{text}", title=title)

    file_name = f"{safe_title}.png"
    path = UPLOADS_ROOT / "artifacts" / file_name
    wrapped = textwrap.wrap(text or title, width=42)[:48]
    if not wrapped:
        wrapped = [title]
    width = 1200
    line_height = 28
    height = max(240, 60 + line_height * len(wrapped))
    image = Image.new("RGB", (width, height), color="#ffffff")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()
    y = 30
    for line in wrapped:
        draw.text((40, y), line, fill="#111827", font=font)
        y += line_height
    image.save(path)
    return file_name, "image/png", attachment_public_url(path.relative_to(PROJECT_ROOT))
